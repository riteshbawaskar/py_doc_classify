"""
Script to generate dummy documents in multiple formats:
- Images: Aadhaar, Passport, PAN Card
- PDF: Experience Letter, Contract
- DOCX: Resume
"""
import os
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create output directory
os.makedirs("dummy_documents", exist_ok=True)

# ===== IMAGE GENERATION =====

def create_text_image(filename, lines, width=800, height=600):
    """Create an image with text"""
    img = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(img)
    
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
        font_bold = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
    except:
        font = ImageFont.load_default()
        font_bold = ImageFont.load_default()
    
    y_position = 30
    for i, line in enumerate(lines):
        if i == 0:  # First line (title) in bold
            draw.text((50, y_position), line, fill='black', font=font_bold)
        else:
            draw.text((50, y_position), line, fill='black', font=font)
        y_position += 35
    
    filepath = os.path.join("dummy_documents", filename)
    img.save(filepath)
    print(f"Created: {filepath}")

# Generate Aadhaar Card Image
aadhaar_lines = [
    "GOVERNMENT OF INDIA",
    "AADHAAR CARD",
    "",
    "Name: Rajesh Kumar",
    "Date of Birth: 15/08/1990",
    "Gender: Male",
    "Aadhaar Number: 1234 5678 9012",
    "",
    "Address:",
    "House No. 45, Sector 12",
    "Pimpri, Maharashtra",
    "PIN: 411018"
]
create_text_image("aadhaar_card.png", aadhaar_lines)

# Generate Passport Image
passport_lines = [
    "REPUBLIC OF INDIA",
    "PASSPORT",
    "",
    "Passport No: K1234567",
    "Surname: SHARMA",
    "Given Names: PRIYA",
    "Nationality: INDIAN",
    "Date of Birth: 22/03/1988",
    "Place of Birth: MUMBAI",
    "Sex: Female",
    "Date of Issue: 10/01/2020",
    "Date of Expiry: 09/01/2030",
    "Type: P",
    "Country Code: IND"
]
create_text_image("passport.jpg", passport_lines)

# Generate PAN Card Image
pan_lines = [
    "INCOME TAX DEPARTMENT",
    "GOVERNMENT OF INDIA",
    "",
    "PERMANENT ACCOUNT NUMBER CARD",
    "",
    "Name: AMIT VERMA",
    "Father's Name: SURESH VERMA",
    "Date of Birth: 05/11/1985",
    "PAN: ABCDE1234F",
    "",
    "This is a permanent account number card",
    "issued by the Income Tax Department."
]
create_text_image("pan_card.png", pan_lines, width=700, height=450)

# ===== PDF GENERATION =====

def create_experience_letter_pdf():
    """Create Experience Letter PDF"""
    filepath = os.path.join("dummy_documents", "experience_letter.pdf")
    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(inch, height - inch, "EXPERIENCE CERTIFICATE")
    
    # Date
    c.setFont("Helvetica", 11)
    c.drawString(inch, height - 1.5*inch, "Date: October 30, 2025")
    
    # Content
    y_position = height - 2*inch
    content = [
        "To Whom It May Concern,",
        "",
        "This is to certify that Mr. Vikram Singh was employed with our organization,",
        "Global Tech Solutions Pvt Ltd, from January 15, 2022 to October 28, 2025.",
        "",
        "During his tenure with us, he worked as a Senior Software Developer in the",
        "Engineering Department. He was responsible for developing and maintaining",
        "enterprise applications, leading development teams, and ensuring code quality.",
        "",
        "Mr. Singh has demonstrated excellent technical skills, professionalism, and",
        "dedication throughout his employment. He has been a valuable team member",
        "and has contributed significantly to our projects.",
        "",
        "We wish him all the best in his future endeavors.",
        "",
        "Regards,",
        "",
        "_____________________",
        "HR Manager",
        "Global Tech Solutions Pvt Ltd",
        "Pune, Maharashtra",
        "hr@globaltechsolutions.com",
    ]
    
    for line in content:
        c.drawString(inch, y_position, line)
        y_position -= 0.3*inch
    
    c.save()
    print(f"Created: {filepath}")

def create_contract_pdf():
    """Create Contract PDF"""
    filepath = os.path.join("dummy_documents", "contract.pdf")
    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(inch, height - inch, "SERVICE AGREEMENT CONTRACT")
    
    # Content
    y_position = height - 1.5*inch
    c.setFont("Helvetica", 10)
    
    content = [
        "This Service Agreement is entered into on November 1, 2025",
        "",
        "BETWEEN:",
        "Service Provider: ABC Consulting Pvt Ltd",
        "Registration No: U74900MH2020PTC123456",
        "Address: 123 Business Park, Mumbai",
        "",
        "AND:",
        "Client: XYZ Technologies Ltd",
        "Registration No: U72200KA2018PTC234567",
        "Address: 456 Tech Hub, Bangalore",
        "",
        "TERMS AND CONDITIONS",
        "",
        "1. SCOPE OF SERVICES",
        "   The Service Provider agrees to provide software development services.",
        "",
        "2. PAYMENT TERMS",
        "   Total Contract Value: INR 10,00,000",
        "   Payment Schedule: Monthly installments",
        "",
        "3. DURATION",
        "   This agreement shall be valid for 12 months from signing date.",
        "",
        "4. CONFIDENTIALITY",
        "   Both parties agree to maintain confidentiality of proprietary information.",
        "",
        "5. TERMINATION",
        "   Either party may terminate with 30 days written notice.",
        "",
        "",
        "_____________________              _____________________",
        "Service Provider                   Client",
        "Date: November 1, 2025             Date: November 1, 2025"
    ]
    
    for line in content:
        if y_position < inch:
            c.showPage()
            y_position = height - inch
            c.setFont("Helvetica", 10)
        c.drawString(inch, y_position, line)
        y_position -= 0.25*inch
    
    c.save()
    print(f"Created: {filepath}")

create_experience_letter_pdf()
create_contract_pdf()

# ===== WORD DOCUMENT GENERATION =====

def create_resume_docx():
    """Create Resume DOCX"""
    doc = Document()
    
    # Title
    title = doc.add_heading('RESUME', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Personal Info
    doc.add_paragraph('Name: Sneha Patel')
    doc.add_paragraph('Email: sneha.patel@email.com')
    doc.add_paragraph('Phone: +91-9876543210')
    doc.add_paragraph('Location: Pune, Maharashtra')
    doc.add_paragraph()
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', 1)
    doc.add_paragraph(
        'Software Engineer with 5 years of experience in web development '
        'and cloud technologies.'
    )
    doc.add_paragraph()
    
    # Work Experience
    doc.add_heading('WORK EXPERIENCE', 1)
    
    doc.add_heading('Senior Software Engineer', 2)
    doc.add_paragraph('Tech Solutions Pvt Ltd, Pune')
    doc.add_paragraph('June 2021 - Present')
    doc.add_paragraph('• Developed scalable web applications using React and Node.js')
    doc.add_paragraph('• Led a team of 3 developers')
    doc.add_paragraph()
    
    doc.add_heading('Software Engineer', 2)
    doc.add_paragraph('Digital Systems, Mumbai')
    doc.add_paragraph('July 2019 - May 2021')
    doc.add_paragraph('• Built REST APIs and microservices')
    doc.add_paragraph()
    
    # Education
    doc.add_heading('EDUCATION', 1)
    doc.add_paragraph('Bachelor of Engineering in Computer Science')
    doc.add_paragraph('University of Pune, 2019')
    doc.add_paragraph()
    
    # Skills
    doc.add_heading('SKILLS', 1)
    doc.add_paragraph('• Programming: Python, JavaScript, Java')
    doc.add_paragraph('• Frameworks: React, Node.js, Django')
    doc.add_paragraph('• Databases: MySQL, MongoDB')
    
    filepath = os.path.join("dummy_documents", "resume.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")

create_resume_docx()

print(f"\n✓ Successfully generated all dummy documents in 'dummy_documents' folder")
print(f"\nGenerated files:")
print(f"  Images (PNG/JPG): aadhaar_card.png, passport.jpg, pan_card.png")
print(f"  PDFs: experience_letter.pdf, contract.pdf")
print(f"  Word: resume.docx")